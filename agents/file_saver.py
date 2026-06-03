from pathlib import Path
from datetime import date
from docx import Document
from fpdf import FPDF

def save_files(company, role, cover_letter, study_guide):

    BASE_DIR = Path(__file__).parent.parent

    folder_name = f"{company}_{role}_{date.today()}".replace(" ","_")
    output_dir = BASE_DIR / "output" / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)

#Save Doc
    doc = Document()
    doc.add_heading(f"Cover Letter - {company}", level=1)
    doc.add_paragraph(cover_letter)
    docx_path = output_dir / "cover_letter.docx"
    doc.save(docx_path)

#Save pdf
    pdf_path = output_dir / "cover_letter.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", fname="C:/Windows/Fonts/arial.ttf", uni=True)
    pdf.set_font("DejaVu", size=12)
    pdf.set_margins(20,20,20)
    pdf.multi_cell(0,8, f"Cover Letter - {company}\n\n{cover_letter}")
    pdf.output(str(pdf_path))

    #Save study guide
    study_guide_path = output_dir / "study_guide.md"
    study_guide_path.write_text(study_guide, encoding="utf-8")

    print(f"Files saved to : {output_dir}")
    return str(docx_path), str(pdf_path), str(study_guide_path)